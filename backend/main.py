import os
import re
import json
import shutil
import subprocess
import sys
import tempfile
import uuid
import zipfile
import io
from datetime import datetime, timedelta, timezone
from typing import List, Optional, Dict

from fastapi import FastAPI, Depends, HTTPException, UploadFile, File, Form
from fastapi.responses import FileResponse, StreamingResponse
from sqlalchemy.orm import Session
from docxtpl import DocxTemplate
from docx import Document as DocxDocument
import jinja2

from database import Base, engine, get_db
import models
from schemas import (
    LoginRequest,
    TokenResponse,
    CategoryCreate,
    CategoryUpdate,
    CategoryOut,
    TemplateOut,
    TemplateDetailOut,
    TemplateFieldOut,
    TemplateFieldsUpdateRequest,
    TemplateFieldCreate,
    PackageCreate,
    PackageUpdate,
    PackageOut,
    PackageItemOut,
    CaseCreate,
    CaseUpdate,
    CaseOut,
    CaseDetailOut,
    CaseFieldValueOut,
    CaseDocumentOut,
    GenerateRequest,
    PreviewRequest,
    PreviewResponse,
    CaseDocumentEditRequest,
)
from security import verify_password, create_access_token
from deps import get_current_user, require_admin
from docx_utils import extract_placeholders
from name_utils import build_clean_filters, build_preview_filters

Base.metadata.create_all(bind=engine)

app = FastAPI()

STORAGE_TEMPLATES_DIR = "/var/lex-dossier/storage/templates"
STORAGE_CASES_DIR = "/var/lex-dossier/storage/cases"


@app.get("/")
def root():
    return {"status": "ok", "message": "Лекс.Досье backend работает"}


@app.get("/health")
def health():
    return {"status": "healthy"}


@app.post("/api/auth/login", response_model=TokenResponse)
def login(data: LoginRequest, db: Session = Depends(get_db)):
    user = db.query(models.User).filter(models.User.email == data.email).first()
    if not user or not verify_password(data.password, user.password_hash):
        raise HTTPException(status_code=401, detail="Неверный email или пароль")
    if not user.is_active:
        raise HTTPException(status_code=403, detail="Пользователь деактивирован")
    token = create_access_token({"sub": str(user.id), "role": user.role})
    return TokenResponse(access_token=token, role=user.role, full_name=user.full_name)


@app.get("/api/auth/me")
def read_current_user(current_user: models.User = Depends(get_current_user)):
    return {
        "full_name": current_user.full_name,
        "email": current_user.email,
        "role": current_user.role,
    }


@app.get("/api/categories", response_model=List[CategoryOut])
def list_categories(
    branch: Optional[str] = None,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    query = db.query(models.Category).filter(models.Category.is_active == True)
    if branch:
        query = query.filter(models.Category.branch == branch)
    return query.order_by(models.Category.sort_order).all()


@app.post("/api/categories", response_model=CategoryOut)
def create_category(
    data: CategoryCreate,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(require_admin),
):
    category = models.Category(
        id=uuid.uuid4(),
        name=data.name,
        branch=data.branch,
        parent_id=data.parent_id,
        sort_order=data.sort_order,
        is_universal=data.is_universal,
    )
    db.add(category)
    db.commit()
    db.refresh(category)
    return category


@app.patch("/api/categories/{category_id}", response_model=CategoryOut)
def update_category(
    category_id: uuid.UUID,
    data: CategoryUpdate,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(require_admin),
):
    category = db.query(models.Category).filter(models.Category.id == category_id).first()
    if not category:
        raise HTTPException(status_code=404, detail="Категория не найдена")
    if data.name is not None:
        category.name = data.name
    if data.branch is not None:
        category.branch = data.branch
    if "parent_id" in data.model_fields_set:
        if data.parent_id == category_id:
            raise HTTPException(status_code=400, detail="Категория не может быть родителем самой себя")
        category.parent_id = data.parent_id
    if data.sort_order is not None:
        category.sort_order = data.sort_order
    if data.is_universal is not None:
        category.is_universal = data.is_universal
    db.commit()
    db.refresh(category)
    return category


def _collect_category_and_descendants(db: Session, category_id: uuid.UUID) -> List[uuid.UUID]:
    ids = [category_id]
    children = db.query(models.Category).filter(models.Category.parent_id == category_id).all()
    for child in children:
        ids.extend(_collect_category_and_descendants(db, child.id))
    return ids


@app.delete("/api/categories/{category_id}")
def delete_category(
    category_id: uuid.UUID,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(require_admin),
):
    """Мягкое удаление (is_active=false), каскадом на все дочерние категории.
    Так сохраняются ссылки из уже существующих шаблонов, пакетов и дел —
    они просто перестают предлагаться для новых, но не ломаются."""
    category = db.query(models.Category).filter(models.Category.id == category_id).first()
    if not category:
        raise HTTPException(status_code=404, detail="Категория не найдена")

    ids = _collect_category_and_descendants(db, category_id)
    db.query(models.Category).filter(models.Category.id.in_(ids)).update(
        {"is_active": False}, synchronize_session=False
    )
    db.commit()
    return {"status": "ok", "deactivated": [str(i) for i in ids]}


@app.get("/api/templates", response_model=List[TemplateOut])
def list_templates(
    category_id: Optional[uuid.UUID] = None,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    query = db.query(models.Template)
    if category_id:
        query = query.filter(models.Template.category_id == category_id)
    if current_user.role != "admin":
        query = query.filter(models.Template.status == "published")
    return query.all()


@app.get("/api/templates/{template_id}", response_model=TemplateDetailOut)
def get_template(
    template_id: uuid.UUID,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    template = db.query(models.Template).filter(models.Template.id == template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="Шаблон не найден")
    if current_user.role != "admin" and template.status != "published":
        raise HTTPException(status_code=403, detail="Шаблон недоступен")

    fields = (
        db.query(models.TemplateField)
        .filter(models.TemplateField.template_id == template_id)
        .order_by(models.TemplateField.sort_order)
        .all()
    )
    return TemplateDetailOut(
        id=template.id,
        category_id=template.category_id,
        name=template.name,
        description=template.description,
        status=template.status,
        file_version=template.file_version,
        doc_group=template.doc_group,
        fields=[TemplateFieldOut.model_validate(f) for f in fields],
    )


def _create_template_record(
    db: Session,
    category: models.Category,
    name: str,
    description: Optional[str],
    file: UploadFile,
    current_user: models.User,
) -> models.Template:
    """Общая часть загрузки шаблона: сохранить файл на диск, создать
    запись Template и автоматически найти в нём поля-плейсхолдеры.
    Используется и для обычной загрузки одного шаблона, и для загрузки
    связанной пары вариантов (см. create_linked_template_pair)."""
    if not file.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Файл должен быть в формате .docx")

    doc_group = "service" if category.branch == "service" else "main"
    template_id = uuid.uuid4()
    template_dir = os.path.join(STORAGE_TEMPLATES_DIR, str(template_id))
    os.makedirs(template_dir, exist_ok=True)
    file_path = os.path.join(template_dir, "v1.docx")

    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    template = models.Template(
        id=template_id,
        category_id=category.id,
        name=name,
        description=description,
        source_file_path=file_path,
        file_version=1,
        status="draft",
        doc_group=doc_group,
        created_by=current_user.id,
    )
    db.add(template)
    db.commit()
    db.refresh(template)

    placeholder_keys = extract_placeholders(file_path)
    for index, key in enumerate(placeholder_keys):
        db.add(models.TemplateField(
            id=uuid.uuid4(),
            template_id=template.id,
            field_key=key,
            label=key.replace("_", " ").capitalize(),
            field_type="text",
            is_required=False,
            is_shared=False,
            sort_order=index,
        ))
    db.commit()
    return template


def _template_to_detail(db: Session, template: models.Template) -> TemplateDetailOut:
    fields = (
        db.query(models.TemplateField)
        .filter(models.TemplateField.template_id == template.id)
        .order_by(models.TemplateField.sort_order)
        .all()
    )
    return TemplateDetailOut(
        id=template.id,
        category_id=template.category_id,
        name=template.name,
        description=template.description,
        status=template.status,
        file_version=template.file_version,
        doc_group=template.doc_group,
        variant_group_id=template.variant_group_id,
        applicant_variant=template.applicant_variant,
        fields=[TemplateFieldOut.model_validate(f) for f in fields],
    )


@app.post("/api/templates", response_model=TemplateDetailOut)
def create_template(
    category_id: uuid.UUID = Form(...),
    name: str = Form(...),
    description: Optional[str] = Form(None),
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: models.User = Depends(require_admin),
):
    category = db.query(models.Category).filter(models.Category.id == category_id).first()
    if not category:
        raise HTTPException(status_code=404, detail="Категория не найдена")
    template = _create_template_record(db, category, name, description, file, current_user)
    return _template_to_detail(db, template)


@app.post("/api/templates/linked-pair", response_model=List[TemplateDetailOut])
def create_linked_template_pair(
    category_id: uuid.UUID = Form(...),
    name: str = Form(...),
    description: Optional[str] = Form(None),
    file_serviceman: UploadFile = File(...),
    file_relatives: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: models.User = Depends(require_admin),
):
    """Загружает СРАЗУ ДВА файла как один "связанный документ" с вариантами
    по типу заявителя (только для направления СВО): один файл — от лица
    самого военнослужащего, другой — от лица родственника. Юрист в карточке
    дела видит ОДИН пункт под именем name; система сама подставляет нужный
    файл по типу заявителя дела (см. _resolve_case_templates).

    Оба шаблона получают ОДНО и то же (введённое один раз) название — так
    оно и попадёт в имя скачиваемого файла. Это осознанно: заводить 10
    разных длинных названий вида "Жалоба ... от военнослужащего в связи
    с ранением ..." под каждую жизненную ситуацию неудобно и путает —
    название остаётся простым и коротким, вариативность текста внутри
    документа скрыта от юриста."""
    category = db.query(models.Category).filter(models.Category.id == category_id).first()
    if not category:
        raise HTTPException(status_code=404, detail="Категория не найдена")
    if category.branch != "svo":
        raise HTTPException(status_code=400, detail="Связанные варианты по типу заявителя доступны только для направления СВО")

    tmpl_serviceman = _create_template_record(db, category, name, description, file_serviceman, current_user)
    tmpl_relatives = _create_template_record(db, category, name, description, file_relatives, current_user)

    group_id = uuid.uuid4()
    tmpl_serviceman.variant_group_id = group_id
    tmpl_serviceman.applicant_variant = "serviceman"
    tmpl_relatives.variant_group_id = group_id
    tmpl_relatives.applicant_variant = "relatives"
    db.commit()
    db.refresh(tmpl_serviceman)
    db.refresh(tmpl_relatives)

    return [_template_to_detail(db, tmpl_serviceman), _template_to_detail(db, tmpl_relatives)]


@app.post("/api/templates/{template_id}/publish", response_model=TemplateOut)
def publish_template(
    template_id: uuid.UUID,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(require_admin),
):
    template = db.query(models.Template).filter(models.Template.id == template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="Шаблон не найден")
    template.status = "published"
    db.commit()
    db.refresh(template)
    return template


@app.post("/api/templates/{template_id}/unlink-variant", response_model=TemplateOut)
def unlink_template_variant(
    template_id: uuid.UUID,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(require_admin),
):
    """Отменяет связку вариантов — ОБА шаблона группы (не только этот)
    снова становятся самостоятельными пунктами списка документов."""
    this_tmpl = db.query(models.Template).filter(models.Template.id == template_id).first()
    if not this_tmpl:
        raise HTTPException(status_code=404, detail="Шаблон не найден")
    if this_tmpl.variant_group_id:
        db.query(models.Template).filter(models.Template.variant_group_id == this_tmpl.variant_group_id).update(
            {"variant_group_id": None, "applicant_variant": None}
        )
        db.commit()
    db.refresh(this_tmpl)
    return this_tmpl


@app.delete("/api/templates/{template_id}")
def delete_template(
    template_id: uuid.UUID,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(require_admin),
):
    template = db.query(models.Template).filter(models.Template.id == template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="Шаблон не найден")

    used_in_packages = (
        db.query(models.TemplatePackageItem)
        .filter(models.TemplatePackageItem.template_id == template_id)
        .count()
    )
    used_in_documents = (
        db.query(models.CaseDocument)
        .filter(models.CaseDocument.template_id == template_id)
        .count()
    )
    if used_in_packages or used_in_documents:
        parts = []
        if used_in_packages:
            parts.append(f"используется в пакетах: {used_in_packages}")
        if used_in_documents:
            parts.append(f"по нему уже сгенерированы документы: {used_in_documents}")
        raise HTTPException(
            status_code=400,
            detail="Нельзя удалить шаблон — " + "; ".join(parts) +
            ". Уберите его из пакетов (сгенерированные документы удалить нельзя — это история дел клиентов).",
        )

    db.query(models.TemplateField).filter(models.TemplateField.template_id == template_id).delete()
    db.delete(template)
    db.commit()

    template_dir = os.path.join(STORAGE_TEMPLATES_DIR, str(template_id))
    if os.path.isdir(template_dir):
        try:
            shutil.rmtree(template_dir)
        except OSError:
            pass

    return {"status": "ok"}


@app.post("/api/templates/{template_id}/fields", response_model=TemplateDetailOut)
def update_template_fields(
    template_id: uuid.UUID,
    data: TemplateFieldsUpdateRequest,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(require_admin),
):
    """Массовое обновление карты полей: тип, обязательность, is_shared и
    shared_group_key (по нему объединяются одинаковые по смыслу поля разных
    шаблонов внутри одного дела/пакета)."""
    template = db.query(models.Template).filter(models.Template.id == template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="Шаблон не найден")

    fields_by_id = {
        f.id: f
        for f in db.query(models.TemplateField).filter(models.TemplateField.template_id == template_id).all()
    }
    for upd in data.fields:
        field = fields_by_id.get(upd.id)
        if not field:
            raise HTTPException(status_code=404, detail=f"Поле {upd.id} не найдено в этом шаблоне")
        field.field_key = upd.field_key
        field.label = upd.label
        field.field_type = upd.field_type
        field.is_required = upd.is_required
        field.is_shared = upd.is_shared
        field.shared_group_key = upd.shared_group_key if upd.is_shared else None
    db.commit()

    fields = (
        db.query(models.TemplateField)
        .filter(models.TemplateField.template_id == template_id)
        .order_by(models.TemplateField.sort_order)
        .all()
    )
    return TemplateDetailOut(
        id=template.id,
        category_id=template.category_id,
        name=template.name,
        description=template.description,
        status=template.status,
        file_version=template.file_version,
        doc_group=template.doc_group,
        fields=[TemplateFieldOut.model_validate(f) for f in fields],
    )


@app.post("/api/templates/{template_id}/fields/add", response_model=TemplateDetailOut)
def add_template_field(
    template_id: uuid.UUID,
    data: TemplateFieldCreate,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(require_admin),
):
    """Добавить поле вручную — на случай, если автоматический разбор .docx
    что-то пропустил, либо в шаблон добавили новый плейсхолдер вручную."""
    template = db.query(models.Template).filter(models.Template.id == template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="Шаблон не найден")

    max_sort = (
        db.query(models.TemplateField)
        .filter(models.TemplateField.template_id == template_id)
        .count()
    )
    field = models.TemplateField(
        id=uuid.uuid4(),
        template_id=template_id,
        field_key=data.field_key,
        label=data.label,
        field_type=data.field_type,
        is_required=data.is_required,
        is_shared=data.is_shared,
        shared_group_key=data.shared_group_key if data.is_shared else None,
        sort_order=max_sort,
    )
    db.add(field)
    db.commit()

    fields = (
        db.query(models.TemplateField)
        .filter(models.TemplateField.template_id == template_id)
        .order_by(models.TemplateField.sort_order)
        .all()
    )
    return TemplateDetailOut(
        id=template.id,
        category_id=template.category_id,
        name=template.name,
        description=template.description,
        status=template.status,
        file_version=template.file_version,
        doc_group=template.doc_group,
        fields=[TemplateFieldOut.model_validate(f) for f in fields],
    )


@app.delete("/api/templates/{template_id}/fields/{field_id}", response_model=TemplateDetailOut)
def delete_template_field(
    template_id: uuid.UUID,
    field_id: uuid.UUID,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(require_admin),
):
    field = (
        db.query(models.TemplateField)
        .filter(models.TemplateField.id == field_id, models.TemplateField.template_id == template_id)
        .first()
    )
    if not field:
        raise HTTPException(status_code=404, detail="Поле не найдено")
    db.delete(field)
    db.commit()

    template = db.query(models.Template).filter(models.Template.id == template_id).first()
    fields = (
        db.query(models.TemplateField)
        .filter(models.TemplateField.template_id == template_id)
        .order_by(models.TemplateField.sort_order)
        .all()
    )
    return TemplateDetailOut(
        id=template.id,
        category_id=template.category_id,
        name=template.name,
        description=template.description,
        status=template.status,
        file_version=template.file_version,
        doc_group=template.doc_group,
        fields=[TemplateFieldOut.model_validate(f) for f in fields],
    )


# ---------- Пакеты ----------

def _package_to_out(db: Session, package: models.TemplatePackage) -> PackageOut:
    rows = (
        db.query(models.TemplatePackageItem, models.Template.name)
        .join(models.Template, models.Template.id == models.TemplatePackageItem.template_id)
        .filter(models.TemplatePackageItem.package_id == package.id)
        .order_by(models.TemplatePackageItem.sort_order)
        .all()
    )
    return PackageOut(
        id=package.id,
        category_id=package.category_id,
        name=package.name,
        is_active=package.is_active,
        items=[
            PackageItemOut(template_id=item.template_id, template_name=name, sort_order=item.sort_order)
            for item, name in rows
        ],
    )


@app.get("/api/packages", response_model=List[PackageOut])
def list_packages(
    category_id: Optional[uuid.UUID] = None,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    query = db.query(models.TemplatePackage).filter(models.TemplatePackage.is_active == True)
    if category_id:
        query = query.filter(models.TemplatePackage.category_id == category_id)
    return [_package_to_out(db, p) for p in query.all()]


@app.post("/api/packages", response_model=PackageOut)
def create_package(
    data: PackageCreate,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(require_admin),
):
    category = db.query(models.Category).filter(models.Category.id == data.category_id).first()
    if not category:
        raise HTTPException(status_code=404, detail="Категория не найдена")

    package = models.TemplatePackage(
        id=uuid.uuid4(),
        category_id=data.category_id,
        name=data.name,
        is_active=True,
    )
    db.add(package)
    db.commit()
    db.refresh(package)

    for index, template_id in enumerate(data.template_ids):
        db.add(
            models.TemplatePackageItem(
                id=uuid.uuid4(),
                package_id=package.id,
                template_id=template_id,
                sort_order=index,
            )
        )
    db.commit()
    return _package_to_out(db, package)


@app.patch("/api/packages/{package_id}", response_model=PackageOut)
def update_package(
    package_id: uuid.UUID,
    data: PackageUpdate,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(require_admin),
):
    package = db.query(models.TemplatePackage).filter(models.TemplatePackage.id == package_id).first()
    if not package:
        raise HTTPException(status_code=404, detail="Пакет не найден")

    if data.name is not None:
        package.name = data.name
    if data.is_active is not None:
        package.is_active = data.is_active
    if data.template_ids is not None:
        db.query(models.TemplatePackageItem).filter(
            models.TemplatePackageItem.package_id == package_id
        ).delete()
        for index, template_id in enumerate(data.template_ids):
            db.add(
                models.TemplatePackageItem(
                    id=uuid.uuid4(),
                    package_id=package_id,
                    template_id=template_id,
                    sort_order=index,
                )
            )
    db.commit()
    return _package_to_out(db, package)


@app.delete("/api/packages/{package_id}")
def delete_package(
    package_id: uuid.UUID,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(require_admin),
):
    package = db.query(models.TemplatePackage).filter(models.TemplatePackage.id == package_id).first()
    if not package:
        raise HTTPException(status_code=404, detail="Пакет не найден")

    # У дел, ссылавшихся на этот пакет, просто снимаем ссылку —
    # сами дела и уже сгенерированные документы не трогаем.
    db.query(models.Case).filter(models.Case.package_id == package_id).update(
        {"package_id": None}, synchronize_session=False
    )
    db.query(models.TemplatePackageItem).filter(
        models.TemplatePackageItem.package_id == package_id
    ).delete()
    db.delete(package)
    db.commit()
    return {"status": "ok"}


# ---------- Дела ----------

# Ключ спец-поля "кто заявитель" для направления СВО. Не привязан ни к
# одному конкретному TemplateField — подставляется в контекст рендера
# каждого документа СВО-дела напрямую (см. _svo_applicant_context), а
# внутри .docx-шаблонов используется через {% if %}/{% elif %}, например:
#   {% if тип_заявителя == "жена" %}являюсь супругой{% elif ... %}
APPLICANT_TYPE_FIELD_KEY = "тип_заявителя"
APPLICANT_TYPE_OPTIONS = ["военнослужащий", "жена", "мать", "отец", "брат", "сестра"]


def _get_case_branch(db: Session, case: models.Case) -> Optional[str]:
    category = db.query(models.Category).filter(models.Category.id == case.category_id).first()
    return category.branch if category else None


def _get_case_applicant_type(db: Session, case: models.Case) -> Optional[str]:
    """Значение спец-поля "кто заявитель" для дела (см. APPLICANT_TYPE_FIELD_KEY),
    или None, если не задано / дело не относится к направлению СВО."""
    if _get_case_branch(db, case) != "svo":
        return None
    value = (
        db.query(models.CaseFieldValue)
        .filter(models.CaseFieldValue.case_id == case.id, models.CaseFieldValue.field_key == APPLICANT_TYPE_FIELD_KEY)
        .first()
    )
    return value.value if value else None


def _svo_applicant_context(db: Session, case: models.Case) -> dict:
    """Если дело относится к направлению СВО — возвращает {тип_заявителя: значение}
    для подмешивания в контекст рендера КАЖДОГО документа этого дела, независимо
    от того, объявлено ли это поле как TemplateField у конкретного шаблона.
    Для гражданских/административных дел возвращает пустой словарь — поведение
    для этого направления не меняется."""
    if _get_case_branch(db, case) != "svo":
        return {}
    return {APPLICANT_TYPE_FIELD_KEY: (_get_case_applicant_type(db, case) or "")}


def _wanted_applicant_variant(applicant_type: Optional[str]) -> str:
    """'военнослужащий' -> вариант документа 'serviceman', любой родственник
    (жена/мать/отец/брат/сестра) или ещё не выбрано -> 'relatives' по
    умолчанию (эти варианты текстуально гораздо ближе друг к другу, чем к
    варианту от лица самого военнослужащего)."""
    return "serviceman" if applicant_type == "военнослужащий" else "relatives"


def _resolve_case_templates(db: Session, case: models.Case, templates: List[models.Template]) -> List[models.Template]:
    """Сворачивает пары шаблонов-вариантов (Template.variant_group_id) до
    ОДНОГО представителя на группу — того, что соответствует типу заявителя
    этого дела. Шаблоны без variant_group_id возвращаются как есть, без
    изменений. Список templates может содержать шаблоны из разных
    категорий/групп одновременно — резолвер работает на всём списке сразу."""
    applicant_type = _get_case_applicant_type(db, case)
    wanted = _wanted_applicant_variant(applicant_type)

    by_group: Dict[uuid.UUID, List[models.Template]] = {}
    result = []
    for t in templates:
        if not t.variant_group_id:
            result.append(t)
            continue
        by_group.setdefault(t.variant_group_id, []).append(t)

    for group_id, members in by_group.items():
        match = next((m for m in members if m.applicant_variant == wanted), None)
        # match может быть None, если админ ещё не доделал пару (например,
        # у обоих шаблонов группы применён один и тот же вариант по ошибке) —
        # тогда просто берём первый, чтобы юрист не остался без документа.
        result.append(match or members[0])

    return result


def _add_business_days(start: datetime, n: int) -> datetime:
    """Прибавляет n рабочих дней (пропускает субботу/воскресенье)."""
    result = start
    added = 0
    while added < n:
        result += timedelta(days=1)
        if result.weekday() < 5:  # 0..4 = пн..пт
            added += 1
    return result


def _refresh_case_status(case: models.Case) -> bool:
    """Ленивая (без крон-задачи) авто-архивация: если дело в работе и с
    момента первой генерации прошло 4 рабочих дня — переводит в архив.
    Возвращает True, если статус был изменён (вызывающий код должен
    сделать commit). Статус draft/ready/archived, выставленные вручную,
    не трогает."""
    if case.status == "in_progress" and case.first_generated_at:
        deadline = _add_business_days(case.first_generated_at, 4)
        now = datetime.now(timezone.utc)
        if now >= deadline:
            case.status = "archived"
            return True
    return False


def _case_to_detail(db: Session, case: models.Case) -> CaseDetailOut:
    field_values = (
        db.query(models.CaseFieldValue)
        .filter(models.CaseFieldValue.case_id == case.id)
        .all()
    )
    documents = (
        db.query(models.CaseDocument, models.Template.name)
        .join(models.Template, models.Template.id == models.CaseDocument.template_id)
        .filter(models.CaseDocument.case_id == case.id)
        .order_by(models.CaseDocument.generated_at)
        .all()
    )
    creator = db.query(models.User).filter(models.User.id == case.created_by).first()

    package_template_ids = []
    if case.package_id:
        pkg_items = (
            db.query(models.TemplatePackageItem)
            .filter(models.TemplatePackageItem.package_id == case.package_id)
            .all()
        )
        pkg_template_ids = [item.template_id for item in pkg_items]
        pkg_templates = db.query(models.Template).filter(models.Template.id.in_(pkg_template_ids)).all()
        package_template_ids = [t.id for t in _resolve_case_templates(db, case, pkg_templates)]

    return CaseDetailOut(
        id=case.id,
        client_name=case.client_name,
        category_id=case.category_id,
        package_id=case.package_id,
        status=case.status,
        created_at=case.created_at,
        created_by_name=creator.full_name if creator else None,
        created_by_email=creator.email if creator else None,
        package_template_ids=package_template_ids,
        fields=[CaseFieldValueOut.model_validate(f) for f in field_values],
        documents=[
            CaseDocumentOut(
                id=doc.id,
                template_id=doc.template_id,
                template_name=template_name,
                has_pdf=bool(doc.pdf_file_path),
                generated_at=doc.generated_at,
            )
            for doc, template_name in documents
        ],
    )


@app.get("/api/cases", response_model=List[CaseOut])
def list_cases(
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    # Юрист и админ пока видят все дела (спецификация не разделяет видимость
    # по автору) — если понадобится ограничить юриста только своими делами,
    # здесь нужно добавить .filter(models.Case.created_by == current_user.id)
    rows = (
        db.query(models.Case, models.User.full_name, models.User.email)
        .outerjoin(models.User, models.User.id == models.Case.created_by)
        .order_by(models.Case.created_at.desc())
        .all()
    )
    result = []
    changed = False
    for case, creator_name, creator_email in rows:
        if _refresh_case_status(case):
            changed = True
        item = CaseOut.model_validate(case)
        item.created_by_name = creator_name
        item.created_by_email = creator_email
        result.append(item)
    if changed:
        db.commit()
    return result


@app.post("/api/cases", response_model=CaseOut)
def create_case(
    data: CaseCreate,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    category = db.query(models.Category).filter(models.Category.id == data.category_id).first()
    if not category:
        raise HTTPException(status_code=404, detail="Категория не найдена")

    if data.package_id:
        package = db.query(models.TemplatePackage).filter(models.TemplatePackage.id == data.package_id).first()
        if not package:
            raise HTTPException(status_code=404, detail="Пакет не найден")

    if data.applicant_type and data.applicant_type not in APPLICANT_TYPE_OPTIONS:
        raise HTTPException(status_code=400, detail="Некорректное значение типа заявителя")

    case = models.Case(
        id=uuid.uuid4(),
        client_name=data.client_name,
        category_id=data.category_id,
        package_id=data.package_id,
        status="draft",
        created_by=current_user.id,
    )
    db.add(case)
    db.flush()  # получаем case.id для CaseFieldValue ниже, коммитим всё разом

    # "Кто заявитель" — только для направления СВО, только если реально
    # передано (гражданские дела это поле не трогает, см. _svo_applicant_context).
    if category.branch == "svo" and data.applicant_type:
        db.add(models.CaseFieldValue(
            id=uuid.uuid4(),
            case_id=case.id,
            field_key=APPLICANT_TYPE_FIELD_KEY,
            value=data.applicant_type,
        ))

    db.commit()
    db.refresh(case)
    return case


@app.get("/api/cases/{case_id}", response_model=CaseDetailOut)
def get_case(
    case_id: uuid.UUID,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    case = db.query(models.Case).filter(models.Case.id == case_id).first()
    if not case:
        raise HTTPException(status_code=404, detail="Дело не найдено")
    if _refresh_case_status(case):
        db.commit()
    return _case_to_detail(db, case)


@app.get("/api/cases/{case_id}/available-templates", response_model=List[TemplateOut])
def list_case_available_templates(
    case_id: uuid.UUID,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    """Список опубликованных шаблонов, которые можно выбрать для этого дела:
    своя категория дела + шаблоны из «общих» категорий (is_universal), уже
    с разрешёнными парами вариантов по типу заявителя (см.
    _resolve_case_templates) — если у документа есть варианты
    служащий/родня, юрист увидит только ОДИН подходящий пункт."""
    case = db.query(models.Case).filter(models.Case.id == case_id).first()
    if not case:
        raise HTTPException(status_code=404, detail="Дело не найдено")

    universal_ids = {
        c.id for c in db.query(models.Category).filter(models.Category.is_universal == True).all()  # noqa: E712
    }
    templates = (
        db.query(models.Template)
        .filter(
            models.Template.status == "published",
            (models.Template.category_id == case.category_id) | (models.Template.category_id.in_(universal_ids)),
        )
        .all()
    )
    return _resolve_case_templates(db, case, templates)


@app.patch("/api/cases/{case_id}", response_model=CaseDetailOut)
def update_case(
    case_id: uuid.UUID,
    data: CaseUpdate,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    case = db.query(models.Case).filter(models.Case.id == case_id).first()
    if not case:
        raise HTTPException(status_code=404, detail="Дело не найдено")
    if data.client_name is not None:
        case.client_name = data.client_name
    if data.status is not None:
        # Обычно статус выставляется автоматически (см. generate_documents и
        # _refresh_case_status) — ручная установка через PATCH остаётся для
        # редких случаев (например, админ хочет досрочно заархивировать).
        case.status = data.status
    db.commit()
    db.refresh(case)
    return _case_to_detail(db, case)


@app.delete("/api/cases/{case_id}")
def delete_case(
    case_id: uuid.UUID,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    case = db.query(models.Case).filter(models.Case.id == case_id).first()
    if not case:
        raise HTTPException(status_code=404, detail="Дело не найдено")

    # Пробуем убрать файлы дела с диска целиком — best effort, ошибка
    # удаления не должна мешать удалению записей в базе. Одной папкой
    # удаляются и сгенерированные документы, и "базовые" файлы ручных
    # правок (case_document_edits.docx_file_path), которые лежат там же.
    case_dir = os.path.join(STORAGE_CASES_DIR, str(case_id))
    if os.path.isdir(case_dir):
        shutil.rmtree(case_dir, ignore_errors=True)

    db.query(models.CaseDocument).filter(models.CaseDocument.case_id == case_id).delete()
    db.query(models.CaseDocumentEdit).filter(models.CaseDocumentEdit.case_id == case_id).delete()
    db.query(models.CaseFieldValue).filter(models.CaseFieldValue.case_id == case_id).delete()
    db.delete(case)
    db.commit()
    return {"status": "ok"}


@app.put("/api/cases/{case_id}/fields", response_model=CaseDetailOut)
def update_case_fields(
    case_id: uuid.UUID,
    data: Dict[str, str],
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    """Пакетное сохранение формы: тело запроса — { field_key: значение, ... }."""
    case = db.query(models.Case).filter(models.Case.id == case_id).first()
    if not case:
        raise HTTPException(status_code=404, detail="Дело не найдено")

    for field_key, value in data.items():
        existing = (
            db.query(models.CaseFieldValue)
            .filter(
                models.CaseFieldValue.case_id == case_id,
                models.CaseFieldValue.field_key == field_key,
            )
            .first()
        )
        if existing:
            existing.value = value
        else:
            db.add(
                models.CaseFieldValue(
                    id=uuid.uuid4(),
                    case_id=case_id,
                    field_key=field_key,
                    value=value,
                )
            )
    db.commit()
    return _case_to_detail(db, case)


def _convert_to_pdf(docx_path: str, out_dir: str) -> Optional[str]:
    """Конвертирует docx в pdf через LibreOffice headless. При любой ошибке
    возвращает None — генерация docx при этом всё равно считается успешной.

    Каждому вызову даётся СВОЙ временный профиль LibreOffice
    (-env:UserInstallation). Без этого повторные вызовы soffice в рамках
    одной генерации пакета (несколько документов подряд) конкурируют за
    один и тот же профиль пользователя, конвертация тихо падает по
    таймауту/блокировке — и PDF не появляется вообще ни у одного документа."""
    if shutil.which("soffice") is None:
        print(
            "[pdf-convert] Бинарник soffice не найден в PATH — LibreOffice, похоже, "
            "не установлен на сервере. Установить: sudo apt install libreoffice --no-install-recommends",
            file=sys.stderr,
        )
        return None

    profile_dir = tempfile.mkdtemp(prefix="lo_profile_")
    try:
        subprocess.run(
            [
                "soffice", "--headless", "--norestore",
                f"-env:UserInstallation=file://{profile_dir}",
                "--convert-to", "pdf", "--outdir", out_dir, docx_path,
            ],
            check=True,
            timeout=90,
            capture_output=True,
        )
        base_name = os.path.splitext(os.path.basename(docx_path))[0]
        pdf_path = os.path.join(out_dir, base_name + ".pdf")
        return pdf_path if os.path.exists(pdf_path) else None
    except Exception as e:
        detail = e.stderr.decode("utf-8", "ignore") if isinstance(e, subprocess.CalledProcessError) and e.stderr else str(e)
        print(f"[pdf-convert] Не удалось сконвертировать {docx_path} в PDF: {detail}", file=sys.stderr)
        return None
    finally:
        shutil.rmtree(profile_dir, ignore_errors=True)


def _set_run_font(run, font_name: str = "Times New Roman"):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:cs"), font_name)


def _normalize_font(docx_path: str, font_name: str = "Times New Roman"):
    """Принудительно приводит шрифт всего документа (включая текст,
    подставленный в плейсхолдеры) к единому шрифту — чтобы вставленные
    значения визуально не отличались от остального текста шаблона.
    При любой ошибке молча пропускает — это не должно ломать генерацию."""
    try:
        doc = DocxDocument(docx_path)

        def process(paragraphs):
            for p in paragraphs:
                for run in p.runs:
                    _set_run_font(run, font_name)

        process(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    process(cell.paragraphs)
        for section in doc.sections:
            process(section.header.paragraphs)
            process(section.footer.paragraphs)
        doc.save(docx_path)
    except Exception:
        pass


def _display_value_for_field(field_type: str, value: str) -> str:
    """Значение поля так, как оно должно попасть в текст документа.
    Поле хранится в базе как есть (для дат — ISO yyyy-mm-dd, это формат
    нативного <input type="date">), но в самом документе дата должна быть
    в привычном для юридических текстов виде дд.мм.гггг."""
    if field_type == "date" and value and re.match(r"^\d{4}-\d{2}-\d{2}$", value):
        y, m, d = value.split("-")
        return f"{d}.{m}.{y}"
    return value


def _format_client_short_name(full_name: str) -> str:
    """«Иванов Иван Иванович» -> «Иванов И.И.»"""
    parts = [p for p in (full_name or "").strip().split() if p]
    if not parts:
        return "Клиент"
    surname = parts[0]
    initials = "".join(f"{p[0].upper()}." for p in parts[1:3])
    return f"{surname} {initials}".strip()


def _sanitize_filename(name: str) -> str:
    name = re.sub(r'[\\/:*?"<>|]', "", name).strip()
    return name or "document"


def _build_document_filename(client_name: str, template_name: str, ext: str) -> str:
    return _sanitize_filename(f"{_format_client_short_name(client_name)} {template_name}") + "." + ext


def _extract_paragraph_texts(docx_path: str) -> List[str]:
    """Текст абзацев основного тела документа (без таблиц, колонтитулов —
    для них редактирование текста пока не поддерживается)."""
    doc = DocxDocument(docx_path)
    return [p.text for p in doc.paragraphs]


def _primary_run(p):
    """Run абзаца, который лучше всего представляет его форматирование —
    самый длинный по тексту, а не всегда runs[0] (первый run нередко
    служебный: пустой / с форматированием-остатком от закладки Word)."""
    if not p.runs:
        return None
    return max(p.runs, key=lambda r: len(r.text or ""))


def _set_paragraph_text_preserving_format(p, text: str):
    """Заменяет весь видимый текст абзаца на text, отдавая его "основному"
    run'у (см. _primary_run) и опустошая остальные.

    КРИТИЧНО: p.runs нужно забрать РОВНО ОДИН раз в переменную и дальше
    работать с этим списком. python-docx возвращает НОВЫЙ список НОВЫХ
    объектов-обёрток Run при каждом обращении к p.runs — даже если это тот
    же самый кусок текста в файле. Из-за этого сравнение "extra is not
    style_run" между результатами ДВУХ РАЗНЫХ обращений к p.runs всегда
    было ложным (True для абсолютно любого run, включая тот же самый) —
    именно это стирало вообще весь текст абзаца, а не только "лишние"
    runs. Баг был обнаружен воспроизведением на реальном файле шаблона."""
    runs = p.runs
    if not runs:
        p.add_run(text)
        return
    style_run = max(runs, key=lambda r: len(r.text or ""))
    style_run.text = text
    for extra in runs:
        if extra is not style_run:
            extra.text = ""


def _strip_markers_in_docx(docx_path: str):
    """Финальная чистка перед выдачей документа юристу: убирает служебные
    unicode-метки предпросмотра (⟪⟫/⟦⟧) прямо из текста абзацев уже
    ГОТОВОГО файла — применяется только к документу, полученному из
    ручной правки (см. save_document_edit).

    Принципиально отличается от старого подхода ("применить текст правки
    по индексу к заново отрендеренному документу"): здесь нет сопоставления
    между ДВУМЯ РАЗНЫМИ версиями документа — мы правим текст абзацев прямо
    внутри ОДНОГО И ТОГО ЖЕ файла, поэтому число абзацев физически не может
    разойтись, и документ не может превратиться в пустые страницы."""
    doc = DocxDocument(docx_path)
    changed = False
    for p in doc.paragraphs:
        original = p.text
        cleaned = _strip_preview_markers(original)
        if cleaned != original:
            _set_paragraph_text_preserving_format(p, cleaned)
            changed = True
    if changed:
        doc.save(docx_path)


def _apply_paragraph_texts(docx_path: str, texts: List[str]):
    """Legacy-путь: применяет текст абзацев к документу ПО ИНДЕКСУ — так
    исторически применялись ручные правки, пока не выяснилось, что для
    этого нужно гарантировать полное структурное совпадение с документом,
    на основе которого правки делались (см. save_document_edit — там
    правки теперь применяются сразу к своему собственному "базовому"
    рендеру, и эта функция для НОВЫХ правок больше не нужна). Оставлена
    только для отображения уже сохранённых старых записей — на случай,
    если они существуют в базе с прошлых версий (без docx_file_path)."""
    doc = DocxDocument(docx_path)
    paragraphs = doc.paragraphs
    n_original = len(paragraphs)
    n_new = len(texts)

    for i in range(min(n_original, n_new)):
        _set_paragraph_text_preserving_format(paragraphs[i], texts[i])

    for i in range(n_new, n_original):
        for run in paragraphs[i].runs:
            run.text = ""

    if n_new > n_original and n_original > 0:
        last_p = paragraphs[-1]
        style_run = _primary_run(last_p)
        for i in range(n_original, n_new):
            new_p = doc.add_paragraph()
            new_p.paragraph_format.alignment = last_p.paragraph_format.alignment
            new_p.paragraph_format.left_indent = last_p.paragraph_format.left_indent
            new_p.paragraph_format.first_line_indent = last_p.paragraph_format.first_line_indent
            run = new_p.add_run(texts[i])
            if style_run is not None:
                run.bold = style_run.bold
                run.italic = style_run.italic
                run.font.size = style_run.font.size
                run.font.name = style_run.font.name

    doc.save(docx_path)


def _strip_preview_markers(text: str) -> str:
    """Убирает служебную unicode-разметку предпросмотра из текста перед
    вставкой в ИТОГОВЫЙ (скачиваемый) документ:
      ⟪значение⟫ (подставлено)        -> просто значение
      ⟦не заполнено: label⟧ (пропуск) -> пусто

    Незаполненное поле убирается в пустоту, а не в читаемую пометку вида
    "[не заполнено: ...]" — юрист и так видит эти метки прямо в окне
    редактирования (raw-текст правки специально не зачищается на фронте,
    см. startEditDoc в app.js), так что скрытой потери информации нет.
    А вот в реальном юридическом документе, готовом к подаче, техническая
    пометка "не заполнено" смотрится недопустимо."""
    text = re.sub(r"⟪([^⟫]*)⟫", r"\1", text)
    text = re.sub(r"⟦[^⟧]*⟧", "", text)
    return text


def _get_manual_edit(db: Session, case_id: uuid.UUID, template_id: uuid.UUID):
    return (
        db.query(models.CaseDocumentEdit)
        .filter(models.CaseDocumentEdit.case_id == case_id, models.CaseDocumentEdit.template_id == template_id)
        .first()
    )


def _format_documents_list(db: Session, template_ids: List[uuid.UUID]) -> str:
    """Нумерованный список названий «основных» документов (doc_group='main')
    из переданного набора — для автоматического плейсхолдера типа
    documents_list. Названия — как в библиотеке шаблонов, без фамилии
    клиента (та добавляется только в имя скачиваемого файла)."""
    if not template_ids:
        return "—"
    templates = (
        db.query(models.Template)
        .filter(models.Template.id.in_(template_ids), models.Template.doc_group == "main")
        .all()
    )
    if not templates:
        return "—"
    order = {tid: i for i, tid in enumerate(template_ids)}
    templates.sort(key=lambda t: order.get(t.id, 0))
    return "; ".join(f"{i + 1}. {t.name}" for i, t in enumerate(templates))


@app.post("/api/cases/{case_id}/generate", response_model=List[CaseDocumentOut])
def generate_documents(
    case_id: uuid.UUID,
    data: GenerateRequest,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    case = db.query(models.Case).filter(models.Case.id == case_id).first()
    if not case:
        raise HTTPException(status_code=404, detail="Дело не найдено")

    # Первое нажатие "Сгенерировать документы" переводит дело из черновика
    # в работу и запускает отсчёт 4 рабочих дней до авто-архивации.
    # Повторные нажатия (обновление уже сгенерированного пакета) статус не трогают.
    if case.status == "draft":
        case.status = "in_progress"
        case.first_generated_at = datetime.now(timezone.utc)

    case_field_values = {
        f.field_key: (f.value or "")
        for f in db.query(models.CaseFieldValue).filter(models.CaseFieldValue.case_id == case_id).all()
    }
    documents_list_text = _format_documents_list(db, data.template_ids)
    svo_applicant_context = _svo_applicant_context(db, case)

    case_dir = os.path.join(STORAGE_CASES_DIR, str(case_id))
    os.makedirs(case_dir, exist_ok=True)

    created_docs = []
    for template_id in data.template_ids:
        template = db.query(models.Template).filter(models.Template.id == template_id).first()
        if not template:
            raise HTTPException(status_code=404, detail=f"Шаблон {template_id} не найден")
        if current_user.role != "admin" and template.status != "published":
            raise HTTPException(status_code=403, detail=f"Шаблон {template_id} недоступен")

        template_fields = (
            db.query(models.TemplateField)
            .filter(models.TemplateField.template_id == template_id)
            .all()
        )
        # Каждому плейсхолдеру шаблона (field_key) подставляем значение дела.
        # Если поле общее (is_shared) и у него задан shared_group_key —
        # значение ищем именно по нему (так реализовано объединение
        # одинаковых по смыслу полей между разными шаблонами пакета).
        # Поле с типом documents_list — не вводится вручную, а вычисляется
        # автоматически: перечень «основных» документов, отмеченных в этом
        # же запуске генерации.
        # Если значения нет — подставляем пустую строку, чтобы docxtpl не падал.
        context = {}
        context.update(svo_applicant_context)
        for f in template_fields:
            if f.field_type == "documents_list":
                context[f.field_key] = documents_list_text
                continue
            lookup_key = f.shared_group_key if (f.is_shared and f.shared_group_key) else f.field_key
            context[f.field_key] = _display_value_for_field(f.field_type, case_field_values.get(lookup_key, ""))

        # Если по этому шаблону в деле уже есть сгенерированный документ —
        # это повторная генерация ("Обновить документы"): перезаписываем его
        # файл и запись, а не плодим дубликаты в списке.
        existing_doc = (
            db.query(models.CaseDocument)
            .filter(models.CaseDocument.case_id == case_id, models.CaseDocument.template_id == template_id)
            .first()
        )
        document_id = existing_doc.id if existing_doc else uuid.uuid4()
        docx_path = os.path.join(case_dir, f"{document_id}.docx")

        try:
            doc = DocxTemplate(template.source_file_path)
            jinja_env = jinja2.Environment()
            jinja_env.filters.update(build_clean_filters())
            doc.render(context, jinja_env)
            doc.save(docx_path)
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Не удалось сформировать документ по шаблону «{template.name}»: {e}",
            )

        # Если по этому документу есть ручные правки текста — они уже
        # применены к своему собственному файлу (см. save_document_edit),
        # просто берём его как есть вместо только что отрендеренного из
        # шаблона. Так исключается сама возможность рассинхронизации
        # абзацев между "тем, что видел юрист" и "тем, что генерируется".
        manual_edit = _get_manual_edit(db, case_id, template_id)
        if manual_edit and manual_edit.docx_file_path and os.path.exists(manual_edit.docx_file_path):
            shutil.copyfile(manual_edit.docx_file_path, docx_path)
            _strip_markers_in_docx(docx_path)
        elif manual_edit:
            # Старая запись правки (сохранена до этого обновления, ещё без
            # docx_file_path) — применяем по индексу с проверкой числа
            # абзацев, как раньше. Как только юрист пересохранит такую
            # правку через редактор, у неё появится docx_file_path и она
            # перейдёт на надёжный путь выше.
            try:
                texts = [_strip_preview_markers(t) for t in json.loads(manual_edit.paragraphs_json)]
                current_count = len(_extract_paragraph_texts(docx_path))
                if len(texts) == current_count:
                    _apply_paragraph_texts(docx_path, texts)
                else:
                    print(
                        f"[manual-edit] Пропущено (устаревшая запись без docx_file_path): "
                        f"{current_count} абзацев в документе, {len(texts)} в правке "
                        f"(template_id={template_id}) — переоткройте документ на редактирование "
                        f"и сохраните правки заново.",
                        file=sys.stderr,
                    )
            except Exception as e:
                print(f"[manual-edit] Ошибка применения устаревшей правки: {e}", file=sys.stderr)

        _normalize_font(docx_path)

        pdf_path = _convert_to_pdf(docx_path, case_dir)

        if existing_doc:
            # Старый PDF (если путь отличается — не должен, т.к. document_id
            # тот же, но на всякий случай подчищаем) и сама запись обновляются
            # на месте — id документа не меняется, значит ссылки на него
            # (например, в открытых вкладках фронта) остаются валидными.
            existing_doc.docx_file_path = docx_path
            existing_doc.pdf_file_path = pdf_path
            existing_doc.generated_at = datetime.now(timezone.utc)
            existing_doc.generated_by = current_user.id
            case_document = existing_doc
        else:
            case_document = models.CaseDocument(
                id=document_id,
                case_id=case_id,
                template_id=template_id,
                docx_file_path=docx_path,
                pdf_file_path=pdf_path,
                generated_by=current_user.id,
            )
            db.add(case_document)
        created_docs.append((case_document, template.name))

    db.commit()

    return [
        CaseDocumentOut(
            id=doc.id,
            template_id=doc.template_id,
            template_name=name,
            has_pdf=bool(doc.pdf_file_path),
            generated_at=doc.generated_at,
        )
        for doc, name in created_docs
    ]


class _PreviewFieldValue(str):
    """Строка для контекста ПРЕДПРОСМОТРА документа: выглядит как обычный
    текст с меткой (⟪значение⟫ или ⟦не заполнено: ...⟧) для {{ поле }},
    но в булевом контексте ({% if поле %}) ведёт себя так же, как вело бы
    себя РЕАЛЬНОЕ значение на генерации — то есть True только если поле
    реально заполнено.

    Без этого шаблоны с условными абзацами (например,
    "{% if телефон %}доп. блок про телефон{% endif %}") в превью считали
    бы поле заполненным всегда, ведь строка-обёртка ⟦не заполнено⟧ сама по
    себе непустая. Из-за этого число абзацев, захваченное при открытии
    ручного редактирования, расходилось с реальным числом абзацев на
    генерации (где то же поле оказывалось пустым и абзац пропадал) — и
    именно это ломало скачанный документ при применении ручных правок."""
    def __new__(cls, display_text: str, is_filled: bool):
        obj = str.__new__(cls, display_text)
        obj._is_filled = is_filled
        return obj

    def __bool__(self):
        return self._is_filled


def _build_preview_context(
    db: Session,
    case: models.Case,
    template_fields: List[models.TemplateField],
    values: Dict[str, str],
    selected_template_ids: List[uuid.UUID],
) -> dict:
    """Контекст для рендера документа С метками предпросмотра (⟪⟫/⟦⟧).
    Общий код для live-предпросмотра и для рендера "базового" файла под
    ручную правку — раньше эта логика была продублирована в двух местах и
    успела разъехаться, что и стало одной из причин рассинхронизации."""
    documents_list_text = _format_documents_list(db, selected_template_ids)
    context = {}
    context.update(_svo_applicant_context(db, case))
    for f in template_fields:
        if f.field_type == "documents_list":
            context[f.field_key] = _PreviewFieldValue(f"⟪{documents_list_text}⟫", bool(documents_list_text))
            continue
        lookup_key = f.shared_group_key if (f.is_shared and f.shared_group_key) else f.field_key
        value = values.get(lookup_key)
        display_value = _display_value_for_field(f.field_type, value) if value else None
        display_text = f"⟪{display_value}⟫" if display_value else f"⟦не заполнено: {f.label}⟧"
        context[f.field_key] = _PreviewFieldValue(display_text, bool(display_value))
    return context


@app.post("/api/cases/{case_id}/preview", response_model=PreviewResponse)
def preview_document(
    case_id: uuid.UUID,
    data: PreviewRequest,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    """Предпросмотр документа по абзацам. Если по этому документу уже
    сохранены ручные правки текста — показывает их (это и есть «текущее
    состояние» документа). Иначе строит абзацы заново по текущим значениям
    формы (возможно ещё не сохранённым), подставляя метку-пропуск вместо
    пустых полей. Ничего не пишет в базу и не создаёт файлов в хранилище,
    кроме одного временного файла для чтения абзацев."""
    case = db.query(models.Case).filter(models.Case.id == case_id).first()
    if not case:
        raise HTTPException(status_code=404, detail="Дело не найдено")

    template = db.query(models.Template).filter(models.Template.id == data.template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="Шаблон не найден")
    if current_user.role != "admin" and template.status != "published":
        raise HTTPException(status_code=403, detail="Шаблон недоступен")

    manual_edit = _get_manual_edit(db, case_id, data.template_id)
    if manual_edit:
        return PreviewResponse(paragraphs=json.loads(manual_edit.paragraphs_json), has_manual_edit=True)

    template_fields = (
        db.query(models.TemplateField)
        .filter(models.TemplateField.template_id == data.template_id)
        .all()
    )
    selected_ids = data.selected_template_ids or [data.template_id]
    context = _build_preview_context(db, case, template_fields, data.values, selected_ids)

    tmp_path = None
    try:
        doc = DocxTemplate(template.source_file_path)
        jinja_env = jinja2.Environment()
        jinja_env.filters.update(build_preview_filters())
        doc.render(context, jinja_env)
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_path = tmp.name
        doc.save(tmp_path)
        paragraphs = _extract_paragraph_texts(tmp_path)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Не удалось построить предпросмотр: {e}")
    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.remove(tmp_path)

    return PreviewResponse(paragraphs=paragraphs, has_manual_edit=False)


@app.put("/api/cases/{case_id}/documents/edit", response_model=PreviewResponse)
def save_document_edit(
    case_id: uuid.UUID,
    data: CaseDocumentEditRequest,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    """Сохраняет ручную правку текста документа.

    Рендерит документ ЗАНОВО (с теми же значениями полей, что сейчас видит
    юрист — data.values/selected_template_ids, как в /preview), применяет
    присланный текст абзацев ПРЯМО К ЭТОМУ рендеру и сохраняет получившийся
    .docx на диск — это и есть docx_file_path, источник истины для будущей
    генерации. Абзацы применяются к тому же самому файлу, из которого
    только что взяты, поэтому число абзацев гарантированно совпадает —
    в отличие от старого подхода, где текст правки накладывался на файл,
    отрендеренный заново уже НА ГЕНЕРАЦИИ (потенциально другой рендер,
    другое число абзацев из-за {% if %} — отсюда были пустые страницы)."""
    case = db.query(models.Case).filter(models.Case.id == case_id).first()
    if not case:
        raise HTTPException(status_code=404, detail="Дело не найдено")

    template = db.query(models.Template).filter(models.Template.id == data.template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="Шаблон не найден")

    template_fields = (
        db.query(models.TemplateField)
        .filter(models.TemplateField.template_id == data.template_id)
        .all()
    )
    selected_ids = data.selected_template_ids or [data.template_id]
    context = _build_preview_context(db, case, template_fields, data.values, selected_ids)

    case_dir = os.path.join(STORAGE_CASES_DIR, str(case_id))
    os.makedirs(case_dir, exist_ok=True)
    edit_docx_path = os.path.join(case_dir, f"_edit_{data.template_id}.docx")

    try:
        doc = DocxTemplate(template.source_file_path)
        jinja_env = jinja2.Environment()
        jinja_env.filters.update(build_preview_filters())
        doc.render(context, jinja_env)
        doc.save(edit_docx_path)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Не удалось построить документ для правки: {e}")

    base_paragraph_count = len(_extract_paragraph_texts(edit_docx_path))
    if base_paragraph_count != len(data.paragraphs):
        # Пока юрист редактировал текст, значения полей успели измениться
        # (например, в другой вкладке) настолько, что число абзацев в
        # шаблоне теперь другое ({% if %} на другое условие). Применять
        # правки по индексу в этом случае небезопасно — лучше явно
        # попросить переоткрыть редактор, чем молча испортить документ.
        raise HTTPException(
            status_code=409,
            detail="Данные дела изменились, пока вы редактировали документ. "
                   "Откройте документ на редактирование заново и повторите правки.",
        )

    doc2 = DocxDocument(edit_docx_path)
    for i, p in enumerate(doc2.paragraphs):
        _set_paragraph_text_preserving_format(p, data.paragraphs[i])
    doc2.save(edit_docx_path)

    existing = _get_manual_edit(db, case_id, data.template_id)
    paragraphs_json = json.dumps(data.paragraphs, ensure_ascii=False)
    if existing:
        existing.paragraphs_json = paragraphs_json
        existing.docx_file_path = edit_docx_path
    else:
        db.add(models.CaseDocumentEdit(
            id=uuid.uuid4(),
            case_id=case_id,
            template_id=data.template_id,
            paragraphs_json=paragraphs_json,
            docx_file_path=edit_docx_path,
        ))
    db.commit()
    return PreviewResponse(paragraphs=data.paragraphs, has_manual_edit=True)


@app.delete("/api/cases/{case_id}/documents/edit")
def discard_document_edit(
    case_id: uuid.UUID,
    template_id: uuid.UUID,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    """Отменяет ручные правки — при следующем открытии документ снова
    будет собираться из значений формы."""
    edit = _get_manual_edit(db, case_id, template_id)
    if edit and edit.docx_file_path and os.path.exists(edit.docx_file_path):
        os.remove(edit.docx_file_path)
    db.query(models.CaseDocumentEdit).filter(
        models.CaseDocumentEdit.case_id == case_id,
        models.CaseDocumentEdit.template_id == template_id,
    ).delete()
    db.commit()
    return {"status": "ok"}


@app.get("/api/cases/{case_id}/documents/{document_id}/download")
def download_document(
    case_id: uuid.UUID,
    document_id: uuid.UUID,
    format: str = "docx",
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    row = (
        db.query(models.CaseDocument, models.Case.client_name, models.Template.name)
        .join(models.Case, models.Case.id == models.CaseDocument.case_id)
        .join(models.Template, models.Template.id == models.CaseDocument.template_id)
        .filter(models.CaseDocument.id == document_id, models.CaseDocument.case_id == case_id)
        .first()
    )
    if not row:
        raise HTTPException(status_code=404, detail="Документ не найден")
    document, client_name, template_name = row

    if format == "pdf":
        if not document.pdf_file_path or not os.path.exists(document.pdf_file_path):
            raise HTTPException(status_code=404, detail="PDF-версия недоступна для этого документа")
        return FileResponse(
            document.pdf_file_path,
            media_type="application/pdf",
            filename=_build_document_filename(client_name, template_name, "pdf"),
        )

    if not os.path.exists(document.docx_file_path):
        raise HTTPException(status_code=404, detail="Файл документа не найден на сервере")
    return FileResponse(
        document.docx_file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=_build_document_filename(client_name, template_name, "docx"),
    )


@app.get("/api/cases/{case_id}/download-all")
def download_all_documents(
    case_id: uuid.UUID,
    format: str = "docx",
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    """Пакет всех сгенерированных документов дела архивом .zip.
    format=docx (по умолчанию, кнопка "Скачать пакет в docx") или
    format=pdf (кнопка "Скачать пакет в PDF") — документы без готового PDF
    в этом случае пропускаются."""
    if format not in ("docx", "pdf"):
        raise HTTPException(status_code=400, detail="Неподдерживаемый формат — docx или pdf")

    case = db.query(models.Case).filter(models.Case.id == case_id).first()
    if not case:
        raise HTTPException(status_code=404, detail="Дело не найдено")

    rows = (
        db.query(models.CaseDocument, models.Template.name)
        .join(models.Template, models.Template.id == models.CaseDocument.template_id)
        .filter(models.CaseDocument.case_id == case_id)
        .order_by(models.CaseDocument.generated_at)
        .all()
    )
    if not rows:
        raise HTTPException(status_code=404, detail="У этого дела ещё нет сгенерированных документов")

    buffer = io.BytesIO()
    used_names = set()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for doc, template_name in rows:
            src_path = doc.pdf_file_path if format == "pdf" else doc.docx_file_path
            if not src_path or not os.path.exists(src_path):
                continue
            base_name = _build_document_filename(case.client_name, template_name, format)
            final_name = base_name
            i = 2
            while final_name in used_names:
                final_name = _sanitize_filename(f"{base_name[:-len(format)-1]} ({i})") + "." + format
                i += 1
            used_names.add(final_name)
            zf.write(src_path, arcname=final_name)
    buffer.seek(0)

    if not used_names:
        raise HTTPException(
            status_code=404,
            detail="PDF-версии ещё не готовы ни для одного документа этого дела" if format == "pdf"
            else "Нет файлов для скачивания",
        )

    zip_filename = _sanitize_filename(f"{_format_client_short_name(case.client_name)} - документы ({format})") + ".zip"
    from urllib.parse import quote
    return StreamingResponse(
        buffer,
        media_type="application/zip",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(zip_filename)}"},
    )
